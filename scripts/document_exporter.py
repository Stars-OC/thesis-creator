#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
论文文档导出工具

支持将 Markdown 论文转换为：
- Word 文档（.docx）
- PDF 文档（.pdf）

使用方法：
    python scripts/document_exporter.py --input paper.md --format docx
    python scripts/document_exporter.py --input paper.md --format pdf
    python scripts/document_exporter.py --input paper.md --format both
"""

import re
import sys
import subprocess
from pathlib import Path
from datetime import datetime
from typing import Optional, Tuple

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[警告] python-docx 未安装，Word 导出功能不可用")
    print("安装命令: pip install python-docx")


def set_chinese_font(run, font_name: str = '宋体', font_size: int = 12, bold: bool = False):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    # 设置中文字体
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_paragraph_format(para, first_line_indent: bool = True, line_spacing: float = 1.5):
    """设置段落格式"""
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)


def create_thesis_document() -> 'Document':
    """创建论文 Word 文档（预设格式）"""
    doc = Document()

    # 设置页面边距（符合中国学术论文标准）
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    return doc


def add_title(doc, text: str):
    """添加论文标题（居中、黑体、三号）"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_chinese_font(run, '黑体', 16, bold=True)
    para.paragraph_format.space_after = Pt(12)


def add_heading(doc, text: str, level: int):
    """添加章节标题"""
    para = doc.add_paragraph()

    # 根据级别设置格式
    if level == 1:  # 一级标题（章标题）
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        set_chinese_font(run, '黑体', 14, bold=True)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    elif level == 2:  # 二级标题（节标题）
        run = para.add_run(text)
        set_chinese_font(run, '黑体', 12, bold=True)
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(3)
    elif level == 3:  # 三级标题
        run = para.add_run(text)
        set_chinese_font(run, '黑体', 12, bold=True)
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)
    else:
        run = para.add_run(text)
        set_chinese_font(run, '黑体', 12, bold=True)

    return para


def add_paragraph(doc, text: str, first_line_indent: bool = True):
    """添加正文段落"""
    para = doc.add_paragraph()
    set_paragraph_format(para, first_line_indent)

    # 处理行内格式
    process_inline_formatting(para, text)

    return para


def process_inline_formatting(para, text: str):
    """处理行内格式（加粗、斜体、代码、上标）"""
    # 移除 Markdown 格式标记但保留文本
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # 加粗
    text = re.sub(r'\*(.+?)\*', r'\1', text)  # 斜体
    text = re.sub(r'`(.+?)`', r'\1', text)  # 行内代码

    # 处理上标引用 <sup>[1]</sup> -> [1]
    text = re.sub(r'<sup>\[(\d+)\]</sup>', r'[\1]', text)

    run = para.add_run(text)
    set_chinese_font(run, '宋体', 12)


def add_code_block(doc, code_lines: list, language: str = ''):
    """添加代码块"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.first_line_indent = Cm(0)

    code_text = '\n'.join(code_lines)
    run = para.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)


def add_table(doc, rows: list):
    """添加表格"""
    if not rows or len(rows) < 1:
        return

    # 创建表格
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 填充内容
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                cell.text = cell_text.strip()
                # 设置单元格字体
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_chinese_font(run, '宋体', 10)

    # 表格后添加空行
    doc.add_paragraph()


def add_list_item(doc, text: str, ordered: bool = False):
    """添加列表项"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.74)
    para.paragraph_format.first_line_indent = Cm(0)

    if ordered:
        run = para.add_run(text)
    else:
        run = para.add_run(f"• {text}")

    set_chinese_font(run, '宋体', 12)


def parse_markdown(content: str) -> list:
    """解析 Markdown 内容，返回结构化数据"""
    lines = content.split('\n')
    elements = []

    in_code_block = False
    code_buffer = []
    code_lang = ''
    in_table = False
    table_rows = []

    for line in lines:
        # 代码块处理
        if line.startswith('```'):
            if in_code_block:
                elements.append(('code', code_buffer, code_lang))
                code_buffer = []
                code_lang = ''
                in_code_block = False
            else:
                in_code_block = True
                code_lang = line[3:].strip()
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        # 表格处理
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            # 跳过分隔行
            if not re.match(r'^\|[\s\-:]+\|', line):
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if cells:
                    table_rows.append(cells)
            continue
        else:
            if in_table and table_rows:
                elements.append(('table', table_rows))
                table_rows = []
                in_table = False

        # 标题
        if line.startswith('# ') and not line.startswith('## '):
            elements.append(('title', line[2:].strip()))
        elif line.startswith('## '):
            elements.append(('h1', line[3:].strip()))
        elif line.startswith('### '):
            elements.append(('h2', line[4:].strip()))
        elif line.startswith('#### '):
            elements.append(('h3', line[5:].strip()))

        # 列表
        elif line.startswith('- ') or line.startswith('* '):
            elements.append(('list', line[2:].strip(), False))
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            elements.append(('list', text, True))

        # 空行
        elif not line.strip():
            pass

        # 普通段落
        else:
            elements.append(('para', line))

    # 处理最后的表格
    if in_table and table_rows:
        elements.append(('table', table_rows))

    return elements


def convert_md_to_docx(input_path: str, output_path: str) -> Tuple[bool, str]:
    """
    将 Markdown 转换为 Word 文档

    Returns:
        (success, message)
    """
    if not DOCX_AVAILABLE:
        return False, "python-docx 未安装"

    try:
        print(f"[信息] 正在读取: {input_path}")

        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 解析 Markdown
        elements = parse_markdown(content)

        # 创建文档
        doc = create_thesis_document()

        # 处理各元素
        for elem in elements:
            elem_type = elem[0]

            if elem_type == 'title':
                add_title(doc, elem[1])
            elif elem_type == 'h1':
                add_heading(doc, elem[1], 1)
            elif elem_type == 'h2':
                add_heading(doc, elem[1], 2)
            elif elem_type == 'h3':
                add_heading(doc, elem[1], 3)
            elif elem_type == 'para':
                add_paragraph(doc, elem[1])
            elif elem_type == 'code':
                add_code_block(doc, elem[1], elem[2] if len(elem) > 2 else '')
            elif elem_type == 'table':
                add_table(doc, elem[1])
            elif elem_type == 'list':
                add_list_item(doc, elem[1], elem[2] if len(elem) > 2 else False)

        # 保存文档
        doc.save(output_path)
        print(f"[成功] Word 文档已保存: {output_path}")

        return True, f"Word 文档已保存到 {output_path}"

    except Exception as e:
        return False, f"转换失败: {str(e)}"


def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> Tuple[bool, str]:
    """
    将 Word 文档转换为 PDF

    支持多种转换方式：
    1. Microsoft Word COM（仅 Windows，需要安装 Word）
    2. LibreOffice（跨平台）
    3. docx2pdf（Python 库）
    """
    # 方法 1: 使用 docx2pdf 库（推荐）
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f"[成功] PDF 文档已保存: {pdf_path}")
        return True, f"PDF 文档已保存到 {pdf_path}"
    except ImportError:
        pass
    except Exception as e:
        print(f"[警告] docx2pdf 转换失败: {e}")

    # 方法 2: 使用 LibreOffice
    try:
        result = subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', '--outdir',
             str(Path(pdf_path).parent), docx_path],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0:
            print(f"[成功] PDF 文档已保存: {pdf_path}")
            return True, f"PDF 文档已保存到 {pdf_path}"
    except FileNotFoundError:
        pass
    except Exception as e:
        print(f"[警告] LibreOffice 转换失败: {e}")

    # 方法 3: 使用 Microsoft Word COM（仅 Windows）
    if sys.platform == 'win32':
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(Path(docx_path).absolute()))
            doc.SaveAs(str(Path(pdf_path).absolute()), FileFormat=17)  # 17 = PDF
            doc.Close()
            word.Quit()
            print(f"[成功] PDF 文档已保存: {pdf_path}")
            return True, f"PDF 文档已保存到 {pdf_path}"
        except ImportError:
            pass
        except Exception as e:
            print(f"[警告] Word COM 转换失败: {e}")

    return False, "PDF 转换需要安装以下工具之一：\n  - docx2pdf: pip install docx2pdf\n  - LibreOffice\n  - Microsoft Word（仅 Windows）"


def export_document(input_path: str, output_dir: str, format_type: str = 'both') -> dict:
    """
    导出文档

    Args:
        input_path: 输入的 Markdown 文件路径
        output_dir: 输出目录
        format_type: 'docx', 'pdf', 或 'both'

    Returns:
        结果字典
    """
    input_path = Path(input_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # 生成文件名
    base_name = input_path.stem
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    results = {
        'input': str(input_path),
        'output_dir': str(output_dir),
        'timestamp': timestamp,
        'formats': {}
    }

    # 转换为 Word
    if format_type in ['docx', 'both']:
        docx_path = output_dir / f"{base_name}.docx"
        success, message = convert_md_to_docx(str(input_path), str(docx_path))
        results['formats']['docx'] = {
            'path': str(docx_path),
            'success': success,
            'message': message
        }

    # 转换为 PDF
    if format_type in ['pdf', 'both']:
        # 先确保有 Word 文件
        docx_path = output_dir / f"{base_name}.docx"
        if not docx_path.exists():
            convert_md_to_docx(str(input_path), str(docx_path))

        pdf_path = output_dir / f"{base_name}.pdf"
        success, message = convert_docx_to_pdf(str(docx_path), str(pdf_path))
        results['formats']['pdf'] = {
            'path': str(pdf_path),
            'success': success,
            'message': message
        }

    return results


def print_export_report(results: dict):
    """打印导出报告"""
    print("\n" + "=" * 50)
    print("[文档导出报告]")
    print("=" * 50)
    print(f"输入文件: {results['input']}")
    print(f"输出目录: {results['output_dir']}")
    print(f"导出时间: {results['timestamp']}")
    print("-" * 50)

    for fmt, info in results['formats'].items():
        status = "[成功]" if info['success'] else "[失败]"
        print(f"{fmt.upper()}: {status}")
        print(f"  路径: {info['path']}")
        if not info['success']:
            print(f"  原因: {info['message']}")
        print()

    print("=" * 50)


if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description='论文文档导出工具')
    parser.add_argument('--input', '-i', required=True, help='输入 Markdown 文件路径')
    parser.add_argument('--output', '-o', default='workspace/final/', help='输出目录')
    parser.add_argument('--format', '-f', choices=['docx', 'pdf', 'both'],
                        default='both', help='输出格式')

    args = parser.parse_args()

    results = export_document(args.input, args.output, args.format)
    print_export_report(results)