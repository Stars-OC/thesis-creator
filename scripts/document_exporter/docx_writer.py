import re
from pathlib import Path
from typing import Tuple

from .config import FORMAT_CONFIG
from .markdown import clean_markdown_content, parse_markdown, strip_doi_links
from .preflight import preflight_validate_images

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[警告] python-docx 未安装，Word 导出功能不可用")
    print("安装命令: pip install python-docx")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    print("[警告] Pillow 未安装，图片尺寸自动缩放功能不可用")
    print("安装命令: pip install Pillow")


def set_chinese_font(run, font_name: str = '宋体', font_size: int = 12, bold: bool = False):
    """设置中文字体（支持加粗）"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold

    # 设置中文字体（East Asia）
    r = run._element
    rPr = r.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)

    # 设置 East Asia 字体
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)

    rFonts.set(qn('w:eastAsia'), font_name)

    # 显式设置东亚文字加粗属性 bCs（run.font.bold 不控制 bCs，必须手动设置）
    bCs_elem = rPr.find(qn('w:bCs'))
    if bCs_elem is None:
        bCs_elem = OxmlElement('w:bCs')
        rPr.append(bCs_elem)
    bCs_elem.set(qn('w:val'), '1' if bold else '0')


def set_paragraph_format(para, first_line_indent: bool = True, line_spacing: float = 1.5):
    """设置段落格式"""
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)


def create_thesis_document() -> 'Document':
    """创建论文 Word 文档（预设格式）"""
    doc = Document()

    # 设置页面边距（符合中国学术论文标准）
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    return doc


HEADING_FONT_SIZES = {1: 16, 2: 14, 3: 12, 4: 10.5}

# 不参与自动编号的特殊章节（使用 Heading 1 样式但不显示"第X章"编号）
UNNUMBERED_H1_SECTIONS = ['摘要', 'Abstract', '致谢', '参考文献']

# 不应出现在目录中的重复标签（功能描述、界面截图等）
NO_TOC_LABELS = {'功能描述', '界面截图', '核心代码', '效果分析'}
ABSTRACT_H1_SECTIONS = {'摘要', 'Abstract'}
TOC_FIELD_INSTRUCTION = 'TOC \\o "1-4" \\h \\z \\u'

# 手动编号剥离正则（转换 docx 时移除，让 Word 多级列表接管）
_NUMBERING_STRIP_PATTERNS = {
    1: re.compile(r'^第\d+章\s+'),
    2: re.compile(r'^\d+\.\d+\s+'),
    3: re.compile(r'^\d+\.\d+\.\d+\s+'),
    4: re.compile(r'^\d+\.\d+\.\d+\.\d+\s+'),
}


def strip_manual_numbering(text: str, level: int) -> str:
    """剥离标题中的手动编号，让 Word 多级列表自动编号接管

    示例：
    - '# 第1章 绪论' → '绪论' (Word 自动编号为 第1章 绪论)
    - '## 1.1 研究背景' → '研究背景' (Word 自动编号为 1.1 研究背景)
    - '### 3.1.1 技术可行性' → '技术可行性' (Word 自动编号为 3.1.1 技术可行性)
    """
    pattern = _NUMBERING_STRIP_PATTERNS.get(level)
    if pattern:
        text = pattern.sub('', text)
    return text


def is_abstract_heading(text: str) -> bool:
    return text.strip() in ABSTRACT_H1_SECTIONS


def should_insert_toc_before_heading(text: str, level: int, seen_abstract: bool, toc_inserted: bool) -> bool:
    if toc_inserted or level != 1:
        return False
    if is_abstract_heading(text):
        return False
    return seen_abstract


def add_title(doc, text: str):
    """添加论文标题（居中、黑体、三号）"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_chinese_font(run, '黑体', 16, bold=True)
    para.paragraph_format.space_after = Pt(12)


def add_heading(doc, text: str, level: int):
    """添加章节标题，使用 Word Heading 样式(支持自动目录和导航窗格)

    Markdown → Word Heading 映射：
    # → Heading 1 (章级，居中)
    ## → Heading 2 (节级，左对齐)
    ### → Heading 3 (小节级，左对齐)
    #### → Heading 4 (段级，左对齐，重复标签不出现在目录)
    """
    # 重复标签（功能描述、界面截图等）不出现在目录/导航窗格
    # 设置大纲级别为"正文文本"(9)，保持 Heading 4 字体格式
    if level == 4 and text in NO_TOC_LABELS:
        heading = doc.add_paragraph()
        heading.style = doc.styles['Heading 4']
        run = heading.add_run(text)
        set_chinese_font(run, '黑体', HEADING_FONT_SIZES.get(4, 10.5), bold=False)
        run.font.color.rgb = RGBColor(0, 0, 0)
        heading.paragraph_format.first_line_indent = Cm(0)
        heading.paragraph_format.space_before = Pt(3)
        heading.paragraph_format.space_after = Pt(3)
        # 设置大纲级别为 9（正文文本），不出现在目录/导航窗格
        pPr = heading._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            heading._element.insert(0, pPr)
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is None:
            outlineLvl = OxmlElement('w:outlineLvl')
            pPr.append(outlineLvl)
        outlineLvl.set(qn('w:val'), '9')
        return heading

    heading = doc.add_heading(text, level=level)

    # 黑体不加粗 + 强制黑色（编号和标题文本均使用此字体）
    font_size = HEADING_FONT_SIZES.get(level, 12)
    for run in heading.runs:
        set_chinese_font(run, '黑体', font_size, bold=False)
        run.font.color.rgb = RGBColor(0, 0, 0)

    # 一级标题居中对齐
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

    # 设置段前段后间距
    elif level == 2:
        heading.paragraph_format.first_line_indent = Cm(0)
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after = Pt(3)
    elif level == 3:
        heading.paragraph_format.first_line_indent = Cm(0)
        heading.paragraph_format.space_before = Pt(3)
        heading.paragraph_format.space_after = Pt(3)

    return heading


def add_paragraph(doc, text: str, first_line_indent: bool = True):
    """添加正文段落"""
    para = doc.add_paragraph()
    set_paragraph_format(para, first_line_indent)

    # 处理行内格式
    process_inline_formatting(para, text)

    return para


def process_inline_formatting(para, text: str):
    """处理行内格式（加粗、斜体、代码、上标）"""
    # 处理上标引用 <sup>[1]</sup> -> [1]
    text = re.sub(r'<sup>\[(\d+)\]</sup>', r'[\1]', text)

    # 逐段处理 **加粗**、*斜体*、`代码`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last_end = 0

    for match in pattern.finditer(text):
        if match.start() > last_end:
            run = para.add_run(text[last_end:match.start()])
            set_chinese_font(run, '宋体', 12)

        if match.group(2):  # **加粗**
            run = para.add_run(match.group(2))
            set_chinese_font(run, '宋体', 12, bold=True)
        elif match.group(3):  # *斜体*
            run = para.add_run(match.group(3))
            set_chinese_font(run, '宋体', 12)
            run.font.italic = True
        elif match.group(4):  # `代码`
            run = para.add_run(match.group(4))
            run.font.name = 'Consolas'
            run.font.size = Pt(10.5)

        last_end = match.end()

    if last_end < len(text):
        run = para.add_run(text[last_end:])
        set_chinese_font(run, '宋体', 12)


def add_code_block(doc, code_lines: list, language: str = ''):
    """添加代码块"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.first_line_indent = Cm(0)

    code_text = '\n'.join(code_lines)
    run = para.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)


def add_table(doc, rows: list, caption: str = ''):
    """
    添加三线表

    三线表规范：
    - 顶线：1.5pt 粗线（w:sz="12"）
    - 表头下线：0.75pt 细线（w:sz="6"）
    - 底线：1.5pt 粗线
    - 其余边框：无
    - 表头文字加粗居中

    Args:
        doc: Word 文档对象
        rows: 表格数据（第一行为表头）
        caption: 表名（在表上方显示，格式：表X.X  名称）
    """
    if not rows or len(rows) < 1:
        return

    # 添加表名（表名在表的上方）
    if caption:
        add_table_caption(doc, caption)

    # 创建表格
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 清除默认边框（先移除 Table Grid 样式的边框）
    table.style = 'Table Grid'

    # 填充内容
    table_font = FORMAT_CONFIG.get('table_font', '黑体')

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text.strip()
                # 设置单元格字体
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        if i == 0:
                            # 表头加粗居中
                            set_chinese_font(run, table_font, 10, bold=True)
                        else:
                            set_chinese_font(run, table_font, 10)

    # 应用三线表边框
    _apply_three_line_borders(table)

    # 表格后添加空行
    doc.add_paragraph()


def _apply_three_line_borders(table):
    """
    应用三线表边框样式

    通过 OxmlElement 直接操作 OOXML，为每个单元格设置边框：
    - 第1行（表头）：顶线粗线 + 底线细线
    - 中间行：无边框
    - 最后1行：底线粗线
    """
    tbl = table._tbl

    # 定义边框尺寸（w:sz 单位为 1/8 磅）
    THICK_SZ = "12"   # 1.5pt = 12/8
    THIN_SZ = "6"     # 0.75pt = 6/8
    BORDER_COLOR = "000000"

    for i, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            # 移除已有边框
            existing = tcPr.find(qn('w:tcBorders'))
            if existing is not None:
                tcPr.remove(existing)

            tcBorders = OxmlElement('w:tcBorders')

            # 设置各边框
            for side in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), BORDER_COLOR)

                if i == 0 and side == 'top':
                    # 顶线：粗线
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), THICK_SZ)
                elif i == 0 and side == 'bottom':
                    # 表头下线：细线
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), THIN_SZ)
                elif i == len(table.rows) - 1 and side == 'bottom':
                    # 底线：粗线
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), THICK_SZ)
                elif side in ('left', 'right'):
                    # 左右边框：无
                    pass

                tcBorders.append(border)

            tcPr.append(tcBorders)


def add_table_caption(doc, caption: str):
    """
    添加表名（表名在表的上方）

    格式：居中、宋体、五号（10.5pt）

    Args:
        doc: Word 文档对象
        caption: 表名文本，如 "表4.1  用户信息表结构"
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(3)

    run = para.add_run(caption)
    set_chinese_font(run, '宋体', 10.5)


def add_list_item(doc, text: str, ordered: bool = False):
    """添加列表项"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.74)
    para.paragraph_format.first_line_indent = Cm(0)

    if ordered:
        run = para.add_run(text)
    else:
        run = para.add_run(f"• {text}")

    set_chinese_font(run, '宋体', 12)


def calculate_image_size(image_path: str, max_width_cm: float = 14.0, max_height_cm: float = 12.0) -> Tuple[float, float]:
    """
    计算图片在 Word 文档中的合适尺寸

    Args:
        image_path: 图片路径
        max_width_cm: 最大宽度（厘米），A4 纸有效宽度约 14cm
        max_height_cm: 最大高度（厘米），避免图片过长

    Returns:
        (width_cm, height_cm): 计算后的尺寸
    """
    if not PIL_AVAILABLE:
        # Pillow 不可用时使用默认宽度
        return max_width_cm, None

    try:
        with Image.open(image_path) as img:
            # 获取原始尺寸（像素）
            orig_width_px, orig_height_px = img.size

            # 假设 Word 文档 DPI 为 96（标准屏幕 DPI）
            # 1 英寸 = 2.54 厘米，1 厘米 ≈ 37.8 像素（96 DPI）
            dpi = 96
            px_per_cm = dpi / 2.54

            # 转换为厘米
            orig_width_cm = orig_width_px / px_per_cm
            orig_height_cm = orig_height_px / px_per_cm

            # 计算缩放比例
            scale_w = max_width_cm / orig_width_cm if orig_width_cm > max_width_cm else 1.0
            scale_h = max_height_cm / orig_height_cm if orig_height_cm > max_height_cm else 1.0

            # 使用较小的缩放比例，确保图片不会超出边界
            scale = min(scale_w, scale_h)

            final_width_cm = orig_width_cm * scale
            final_height_cm = orig_height_cm * scale

            return final_width_cm, final_height_cm

    except Exception as e:
        print(f"[警告] 无法读取图片尺寸: {e}")
        return max_width_cm, None


def add_image(doc, image_path: str, width_cm: float = None, max_width_cm: float = 14.0, max_height_cm: float = 12.0, base_dir: str = ''):
    """
    添加图片到文档（自动缩放）

    修复说明：python-docx 中图片必须通过 doc.add_picture() 添加，
    它会自动创建一个包含图片的新段落。我们再设置该段落居中对齐。

    Args:
        doc: Word 文档对象
        image_path: 图片路径（相对路径或绝对路径）
        width_cm: 图片宽度（厘米），如果为 None 则自动计算
        max_width_cm: 最大宽度限制（厘米）
        max_height_cm: 最大高度限制（厘米）
        base_dir: 基础目录，用于解析相对路径

    Returns:
        bool: 是否成功添加图片
    """
    from pathlib import Path

    # 处理路径
    if base_dir and not Path(image_path).is_absolute():
        full_path = Path(base_dir) / image_path
    else:
        full_path = Path(image_path)

    # 规范化路径（处理 Windows 反斜杠问题）
    full_path = full_path.resolve()

    # 检查文件是否存在
    if not full_path.exists():
        print(f"[警告] 图片文件不存在: {full_path}")
        return False

    # 检查文件扩展名是否为支持的图片格式
    supported_formats = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif', '.emf', '.wmf'}
    if full_path.suffix.lower() not in supported_formats:
        print(f"[警告] 不支持的图片格式: {full_path.suffix} (文件: {full_path})")
        return False

    try:
        # 计算合适的尺寸
        if width_cm is None:
            calc_width, calc_height = calculate_image_size(str(full_path), max_width_cm, max_height_cm)
            width_cm = calc_width
        else:
            width_cm = min(width_cm, max_width_cm)

        # 关键修复：使用 doc.add_picture() 直接添加图片
        # 这会自动创建一个包含图片的新段落
        # 之前的 run.add_picture() 不是 python-docx 的标准用法，导致图片不显示
        doc.add_picture(str(full_path), width=Cm(width_cm))

        # 获取刚添加的图片段落（最后一个段落），设置居中对齐
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_paragraph.paragraph_format.first_line_indent = Cm(0)
        last_paragraph.paragraph_format.space_before = Pt(6)
        last_paragraph.paragraph_format.space_after = Pt(3)

        print(f"[成功] 图片已插入: {full_path.name} (宽度: {width_cm:.1f}cm)")
        return True
    except Exception as e:
        print(f"[警告] 插入图片失败: {full_path} - 错误: {e}")

        # 失败时在文档中插入占位文字，方便用户手动补图
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f"[图片缺失: {full_path.name}]")
            set_chinese_font(run, '宋体', 10.5)
            run.font.color.rgb = RGBColor(255, 0, 0)  # 红色标记
        except:
            pass

        return False


def add_figure_caption(doc, caption: str):
    """
    添加图注（图片说明）

    格式：居中、宋体、五号（10.5pt）

    Args:
        doc: Word 文档对象
        caption: 图注文本，如 "图4-1 系统架构图"
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(12)

    run = para.add_run(caption)
    set_chinese_font(run, '宋体', 10.5)  # 五号字体


def add_page_break(doc):
    """
    添加分页符

    在当前位置插入分页符，后续内容从新页面开始。

    Args:
        doc: Word 文档对象
    """
    # 使用 docx 的 add_page_break 方法
    doc.add_page_break()


def enable_update_fields_on_open(doc):
    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        settings.append(update_fields)
    update_fields.set(qn('w:val'), 'true')


def add_table_of_contents(doc):
    enable_update_fields_on_open(doc)

    heading_para = doc.add_paragraph()
    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_para.paragraph_format.first_line_indent = Cm(0)

    heading_run = heading_para.add_run('目录')
    set_chinese_font(heading_run, '黑体', 14, bold=True)
    heading_run.font.color.rgb = RGBColor(0, 0, 0)

    toc_para = doc.add_paragraph()
    toc_para.paragraph_format.first_line_indent = Cm(0)

    begin_run = toc_para.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    fld_begin.set(qn('w:dirty'), 'true')
    begin_run._r.append(fld_begin)

    instr_run = toc_para.add_run()
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = TOC_FIELD_INSTRUCTION
    instr_run._r.append(instr_text)

    separate_run = toc_para.add_run()
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    separate_run._r.append(fld_separate)

    placeholder_run = toc_para.add_run('右键更新目录')
    set_chinese_font(placeholder_run, '宋体', 10.5)

    end_run = toc_para.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    end_run._r.append(fld_end)

    add_page_break(doc)


def setup_heading_numbering(doc):
    """设置论文标题多级列表自动编号

    编号格式：
    - Heading 1: 第1章, 第2章, ...
    - Heading 2: 1.1, 1.2, ...
    - Heading 3: 1.1.1, 1.2.1, ...
    - Heading 4: 1.1.1.1, ...

    摘要/Abstract/致谢/参考文献 使用 Heading 1 样式但不编号（numId=0）
    """
    try:
        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part.element
    except (KeyError, AttributeError):
        print("[警告] 无法访问 numbering part，标题编号功能不可用")
        return None

    # 查找已存在的 abstractNum ID
    existing_ids = [
        int(an.get(qn('w:abstractNumId')))
        for an in numbering_elm.findall(qn('w:abstractNum'))
        if an.get(qn('w:abstractNumId')) is not None
    ]
    new_abstract_num_id = max(existing_ids) + 1 if existing_ids else 0

    # 创建 abstractNum
    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), str(new_abstract_num_id))

    multi_level_type = OxmlElement('w:multiLevelType')
    multi_level_type.set(qn('w:val'), 'multilevel')
    abstract_num.append(multi_level_type)

    # 定义各编号级别
    levels = [
        (0, 'Heading1', 'decimal', '第%1章'),
        (1, 'Heading2', 'decimal', '%1.%2'),
        (2, 'Heading3', 'decimal', '%1.%2.%3'),
        (3, 'Heading4', 'decimal', '%1.%2.%3.%4'),
    ]

    for ilvl, pStyle, numFmt, lvlText in levels:
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(ilvl))

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        numFmt_elem = OxmlElement('w:numFmt')
        numFmt_elem.set(qn('w:val'), numFmt)
        lvl.append(numFmt_elem)

        pStyle_elem = OxmlElement('w:pStyle')
        pStyle_elem.set(qn('w:val'), pStyle)
        lvl.append(pStyle_elem)

        lvlText_elem = OxmlElement('w:lvlText')
        lvlText_elem.set(qn('w:val'), lvlText)
        lvl.append(lvlText_elem)

        lvlJc_elem = OxmlElement('w:lvlJc')
        lvlJc_elem.set(qn('w:val'), 'left')
        lvl.append(lvlJc_elem)

        suff_elem = OxmlElement('w:suff')
        suff_elem.set(qn('w:val'), 'space')
        lvl.append(suff_elem)

        abstract_num.append(lvl)

    # abstractNum 必须在 num 元素之前
    first_num = numbering_elm.find(qn('w:num'))
    if first_num is not None:
        first_num.addprevious(abstract_num)
    else:
        numbering_elm.insert(0, abstract_num)

    # 创建 num 引用
    existing_num_ids = [
        int(n.get(qn('w:numId')))
        for n in numbering_elm.findall(qn('w:num'))
        if n.get(qn('w:numId')) is not None
    ]
    new_num_id = max(existing_num_ids) + 1 if existing_num_ids else 1

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(new_num_id))
    abstract_num_id_ref = OxmlElement('w:abstractNumId')
    abstract_num_id_ref.set(qn('w:val'), str(new_abstract_num_id))
    num.append(abstract_num_id_ref)
    numbering_elm.append(num)

    return new_num_id


def apply_numbering_to_headings(doc, num_id):
    """为标题段落应用编号

    正常章节自动编号，摘要/Abstract/致谢/参考文献 禁用编号（numId=0）
    """
    if num_id is None:
        return

    for para in doc.paragraphs:
        style_name = para.style.name
        if not style_name or not style_name.startswith('Heading'):
            continue

        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            para._element.insert(0, pPr)

        # 检查是否为不编号的特殊章节
        is_unnumbered = (
            style_name == 'Heading 1' and
            any(section in para.text for section in UNNUMBERED_H1_SECTIONS)
        )

        # 移除已有的 numId
        existing_numId = pPr.find(qn('w:numId'))
        if existing_numId is not None:
            pPr.remove(existing_numId)

        numId_elem = OxmlElement('w:numId')
        if is_unnumbered:
            numId_elem.set(qn('w:val'), '0')
        else:
            numId_elem.set(qn('w:val'), str(num_id))
        pPr.append(numId_elem)


def convert_md_to_docx(input_path: str, output_path: str) -> Tuple[bool, str]:
    """
    将 Markdown 转换为 Word 文档

    Returns:
        (success, message)
    """
    if not DOCX_AVAILABLE:
        return False, "python-docx 未安装"

    try:
        print(f"[信息] 正在读取: {input_path}")

        # 获取输入文件所在目录，用于解析图片相对路径
        input_file = Path(input_path)
        base_dir = input_file.parent

        preflight_ok, preflight_message = preflight_validate_images(input_file)
        if not preflight_ok:
            return False, preflight_message

        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 导出 docx 时清除 DOI 链接（Markdown 终稿中保留 DOI 便于溯源）
        content = strip_doi_links(content)

        # 清理无意义字符
        content = clean_markdown_content(content)

        # 解析 Markdown
        elements = parse_markdown(content)

        # 创建文档
        doc = create_thesis_document()

        # 统计信息
        image_count = 0
        image_failed = []
        seen_abstract = False
        toc_inserted = False

        # 处理各元素
        for elem in elements:
            elem_type = elem[0]

            if elem_type == 'title':
                add_title(doc, elem[1])
            elif elem_type in {'h1', 'h2', 'h3', 'h4'}:
                level = int(elem_type[1])
                heading_text = elem[1]
                if should_insert_toc_before_heading(heading_text, level, seen_abstract, toc_inserted):
                    add_table_of_contents(doc)
                    toc_inserted = True
                add_heading(doc, heading_text, level)
                if level == 1 and is_abstract_heading(heading_text):
                    seen_abstract = True
            elif elem_type == 'para':
                add_paragraph(doc, elem[1])
            elif elem_type == 'code':
                add_code_block(doc, elem[1], elem[2] if len(elem) > 2 else '')
            elif elem_type == 'table':
                add_table(doc, elem[1])
            elif elem_type == 'list':
                add_list_item(doc, elem[1], elem[2] if len(elem) > 2 else False)
            elif elem_type == 'pagebreak':
                # 分页符处理
                add_page_break(doc)
            elif elem_type == 'image':
                # 处理图片（自动缩放）
                img_path = elem[1]
                alt_text = elem[2] if len(elem) > 2 else ''
                success = add_image(
                    doc,
                    img_path,
                    width_cm=float(FORMAT_CONFIG.get("image_width", 12)),
                    max_width_cm=float(FORMAT_CONFIG.get("max_image_width", 14)),
                    max_height_cm=float(FORMAT_CONFIG.get("max_image_height", 18)),
                    base_dir=str(base_dir),
                )
                if success:
                    image_count += 1
                    # 添加图注（如果有说明文字）
                    if alt_text:
                        add_figure_caption(doc, alt_text)
                else:
                    image_failed.append(img_path)

        # 保存文档
        doc.save(output_path)
        print(f"[成功] Word 文档已保存: {output_path}")

        # 输出图片统计
        if image_count > 0:
            print(f"[信息] 成功插入 {image_count} 张图片")
        if image_failed:
            print(f"[警告] {len(image_failed)} 张图片插入失败: {image_failed}")

        return True, f"Word 文档已保存到 {output_path}"

    except Exception as e:
        return False, f"转换失败: {str(e)}"
