#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown 转 Word 文档转换器（独立版）
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEADING_FONT_SIZES = {
    1: 16,  # Heading 1 → 黑体 16pt (章级标题)
    2: 14,  # Heading 2 → 黑体 14pt (节级标题)
    3: 12,  # Heading 3 → 黑体 12pt (小节级标题)
    4: 10.5,  # Heading 4 → 黑体 10.5pt (段级标题)
}

# 不参与自动编号的特殊章节
UNNUMBERED_H1_SECTIONS = ['摘要', 'Abstract', '致谢', '参考文献']

# 不应出现在目录中的重复标签
NO_TOC_LABELS = {'功能描述', '界面截图', '核心代码', '效果分析'}

# 手动编号剥离正则（转换 docx 时移除，让 Word 多级列表接管）
_NUMBERING_STRIP_PATTERNS = {
    1: re.compile(r'^第\d+章\s+'),
    2: re.compile(r'^\d+\.\d+\s+'),
    3: re.compile(r'^\d+\.\d+\.\d+\s+'),
    4: re.compile(r'^\d+\.\d+\.\d+\.\d+\s+'),
}


def strip_manual_numbering(text, level):
    """剥离标题中的手动编号，让 Word 多级列表自动编号接管"""
    pattern = _NUMBERING_STRIP_PATTERNS.get(level)
    if pattern:
        text = pattern.sub('', text)
    return text


def set_chinese_font(run, font_name='宋体', font_size=12, bold=False):
    """设置中文字体（支持加粗控制，含东亚文字 bCs 属性）"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold

    r = run._element
    rPr = r.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

    # 显式设置东亚文字加粗属性 bCs（run.font.bold 不控制 bCs）
    bCs_elem = rPr.find(qn('w:bCs'))
    if bCs_elem is None:
        bCs_elem = OxmlElement('w:bCs')
        rPr.append(bCs_elem)
    bCs_elem.set(qn('w:val'), '1' if bold else '0')


def create_thesis_document():
    """创建论文Word文档"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    return doc


def add_heading(doc, text, level):
    """添加标题（黑体不加粗，强制黑色，使用 Word Heading 样式）
    重复标签（功能描述等）大纲级别设为正文文本，不出现在目录/导航窗格
    """
    if level == 4 and text in NO_TOC_LABELS:
        heading = doc.add_paragraph()
        heading.style = doc.styles['Heading 4']
        run = heading.add_run(text)
        set_chinese_font(run, '黑体', HEADING_FONT_SIZES.get(4, 10.5), bold=False)
        run.font.color.rgb = RGBColor(0, 0, 0)
        heading.paragraph_format.first_line_indent = Cm(0)
        heading.paragraph_format.space_before = Pt(3)
        heading.paragraph_format.space_after = Pt(3)
        pPr = heading._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            heading._element.insert(0, pPr)
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is None:
            outlineLvl = OxmlElement('w:outlineLvl')
            pPr.append(outlineLvl)
        outlineLvl.set(qn('w:val'), '9')
        return heading

    heading = doc.add_heading(text, level=level)
    font_size = HEADING_FONT_SIZES.get(level, 12)
    for run in heading.runs:
        set_chinese_font(run, '黑体', font_size, bold=False)
        run.font.color.rgb = RGBColor(0, 0, 0)

    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    elif level == 2:
        heading.paragraph_format.first_line_indent = Cm(0)
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after = Pt(3)
    elif level == 3:
        heading.paragraph_format.first_line_indent = Cm(0)
        heading.paragraph_format.space_before = Pt(3)
        heading.paragraph_format.space_after = Pt(3)

    return heading


def setup_heading_numbering(doc):
    """设置论文标题多级列表自动编号

    编号格式：
    - Heading 1: 第1章, 第2章, ...
    - Heading 2: 1.1, 1.2, ...
    - Heading 3: 1.1.1, 1.2.1, ...
    - Heading 4: 1.1.1.1, ...
    """
    try:
        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part.element
    except (KeyError, AttributeError):
        print("[警告] 无法访问 numbering part，标题编号功能不可用")
        return None

    existing_ids = [
        int(an.get(qn('w:abstractNumId')))
        for an in numbering_elm.findall(qn('w:abstractNum'))
        if an.get(qn('w:abstractNumId')) is not None
    ]
    new_abstract_num_id = max(existing_ids) + 1 if existing_ids else 0

    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), str(new_abstract_num_id))

    multi_level_type = OxmlElement('w:multiLevelType')
    multi_level_type.set(qn('w:val'), 'multilevel')
    abstract_num.append(multi_level_type)

    levels = [
        (0, 'Heading1', 'decimal', '第%1章'),
        (1, 'Heading2', 'decimal', '%1.%2'),
        (2, 'Heading3', 'decimal', '%1.%2.%3'),
        (3, 'Heading4', 'decimal', '%1.%2.%3.%4'),
    ]

    for ilvl, pStyle, numFmt, lvlText in levels:
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(ilvl))

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        numFmt_elem = OxmlElement('w:numFmt')
        numFmt_elem.set(qn('w:val'), numFmt)
        lvl.append(numFmt_elem)

        pStyle_elem = OxmlElement('w:pStyle')
        pStyle_elem.set(qn('w:val'), pStyle)
        lvl.append(pStyle_elem)

        lvlText_elem = OxmlElement('w:lvlText')
        lvlText_elem.set(qn('w:val'), lvlText)
        lvl.append(lvlText_elem)

        lvlJc_elem = OxmlElement('w:lvlJc')
        lvlJc_elem.set(qn('w:val'), 'left')
        lvl.append(lvlJc_elem)

        suff_elem = OxmlElement('w:suff')
        suff_elem.set(qn('w:val'), 'space')
        lvl.append(suff_elem)

        abstract_num.append(lvl)

    first_num = numbering_elm.find(qn('w:num'))
    if first_num is not None:
        first_num.addprevious(abstract_num)
    else:
        numbering_elm.insert(0, abstract_num)

    existing_num_ids = [
        int(n.get(qn('w:numId')))
        for n in numbering_elm.findall(qn('w:num'))
        if n.get(qn('w:numId')) is not None
    ]
    new_num_id = max(existing_num_ids) + 1 if existing_num_ids else 1

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(new_num_id))
    abstract_num_id_ref = OxmlElement('w:abstractNumId')
    abstract_num_id_ref.set(qn('w:val'), str(new_abstract_num_id))
    num.append(abstract_num_id_ref)
    numbering_elm.append(num)

    return new_num_id


def apply_numbering_to_headings(doc, num_id):
    """为标题段落应用编号（摘要/Abstract/致谢/参考文献 不编号）"""
    if num_id is None:
        return

    for para in doc.paragraphs:
        style_name = para.style.name
        if not style_name or not style_name.startswith('Heading'):
            continue

        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            para._element.insert(0, pPr)

        is_unnumbered = (
            style_name == 'Heading 1' and
            any(section in para.text for section in UNNUMBERED_H1_SECTIONS)
        )

        existing_numId = pPr.find(qn('w:numId'))
        if existing_numId is not None:
            pPr.remove(existing_numId)

        numId_elem = OxmlElement('w:numId')
        if is_unnumbered:
            numId_elem.set(qn('w:val'), '0')
        else:
            numId_elem.set(qn('w:val'), str(num_id))
        pPr.append(numId_elem)


def add_paragraph(doc, text, first_line_indent=True):
    """添加段落"""
    para = doc.add_paragraph()
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5

    run = para.add_run(text)
    set_chinese_font(run, '宋体', 12)
    return para


def parse_markdown_line(line):
    """解析Markdown行"""
    # 标题（较长前缀优先匹配）
    if line.startswith('#### '):
        return ('h4', line[5:])
    elif line.startswith('### '):
        return ('h3', line[4:])
    elif line.startswith('## '):
        return ('h2', line[3:])
    elif line.startswith('# '):
        return ('h1', line[2:])

    # 列表
    if line.startswith('- ') or line.startswith('* '):
        return ('list', line[2:])
    if re.match(r'^\d+\.\s', line):
        return ('list', re.sub(r'^\d+\.\s', '', line))

    # 表格
    if line.startswith('|'):
        return ('table', line)

    # 代码块
    if line.startswith('```'):
        return ('code', line)

    # 普通段落
    if line.strip():
        return ('para', line)

    return ('empty', '')


def process_inline_formatting(doc, para, text):
    """处理行内格式（加粗、斜体、代码）"""
    run = para.add_run(text)
    set_chinese_font(run, '宋体', 12)
    return para


def convert_md_to_docx(input_path, output_path):
    """将Markdown转换为Word文档"""
    print(f"正在读取: {input_path}")

    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    doc = create_thesis_document()

    in_code_block = False
    code_buffer = []
    in_table = False
    table_rows = []

    for line in lines:
        parsed = parse_markdown_line(line)
        line_type, line_content = parsed

        # 代码块处理
        if line_type == 'code':
            if in_code_block:
                if code_buffer:
                    code_text = '\n'.join(code_buffer)
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Cm(1)
                    run = para.add_run(code_text)
                    set_chinese_font(run, 'Consolas', 10)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        # 表格处理
        if line_type == 'table':
            if not in_table:
                in_table = True
                table_rows = []
            if not re.match(r'^\|[\s\-:]+\|', line):
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if cells:
                    table_rows.append(cells)
            continue
        else:
            if in_table and table_rows:
                if len(table_rows) > 0:
                    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    table.style = 'Table Grid'
                    for i, row_data in enumerate(table_rows):
                        for j, cell_text in enumerate(row_data):
                            if j < len(table.rows[i].cells):
                                table.rows[i].cells[j].text = cell_text
                table_rows = []
                in_table = False

        # 标题处理
        if line_type == 'h1':
            add_heading(doc, line_content, 1)
        elif line_type == 'h2':
            add_heading(doc, line_content, 2)
        elif line_type == 'h3':
            add_heading(doc, line_content, 3)
        elif line_type == 'h4':
            add_heading(doc, line_content, 4)
        elif line_type == 'list':
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(line_content)
            set_chinese_font(run, '宋体', 12)
        elif line_type == 'para':
            text = line_content
            text = re.sub(r'<sup>\[(\d+)\]</sup>', r'[\1]', text)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            add_paragraph(doc, text)
        elif line_type == 'empty':
            pass

    # 处理最后的表格
    if in_table and table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table.style = 'Table Grid'
        for i, row_data in enumerate(table_rows):
            for j, cell_text in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    table.rows[i].cells[j].text = cell_text

    # 保存文档
    doc.save(output_path)
    print(f"Word文档已保存到: {output_path}")
    return output_path


if __name__ == '__main__':
    import sys
    input_file = sys.argv[1] if len(sys.argv) > 1 else "workspace/final/论文终稿.md"
    output_file = sys.argv[2] if len(sys.argv) > 2 else "workspace/final/论文终稿.docx"

    convert_md_to_docx(input_file, output_file)